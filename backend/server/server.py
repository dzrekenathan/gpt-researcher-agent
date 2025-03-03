import json
import os
import shutil
from typing import Dict, List


from fastapi import HTTPException, Response
from fastapi.responses import JSONResponse
import os
import io
from docx import Document
import boto3
from botocore.exceptions import ClientError
import logging

from fastapi import FastAPI, Request, WebSocket, WebSocketDisconnect, File, UploadFile, Header, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from gpt_researcher.utils.enum import ReportType, Tone, ReportSource
from backend.server.websocket_manager import WebSocketManager, run_agent
from backend.server.server_utils import (
    get_config_dict,
    update_environment_variables, handle_file_upload, handle_file_deletion,
    execute_multi_agents, handle_websocket_communication
)
from fastapi import Response, HTTPException
from docx import Document
import io
import os


from gpt_researcher.utils.logging_config import setup_research_logging

import logging

# Get logger instance
logger = logging.getLogger(__name__)

# Don't override parent logger settings
logger.propagate = True

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.StreamHandler()  # Only log to console
    ]
)

# Models

# Pydantic model for the new endpoint
class GenerateReportRequest(BaseModel):
    prompt: str
    tone: Tone
    report_source: ReportSource
    report_type: ReportType

class ResearchRequest(BaseModel):
    task: str
    report_type: str
    agent: str


class ConfigRequest(BaseModel):
    ANTHROPIC_API_KEY: str
    TAVILY_API_KEY: str
    LANGCHAIN_TRACING_V2: str
    LANGCHAIN_API_KEY: str
    OPENAI_API_KEY: str
    DOC_PATH: str
    RETRIEVER: str
    GOOGLE_API_KEY: str = ''
    GOOGLE_CX_KEY: str = ''
    BING_API_KEY: str = ''
    SEARCHAPI_API_KEY: str = ''
    SERPAPI_API_KEY: str = ''
    SERPER_API_KEY: str = ''
    SEARX_URL: str = ''
    XAI_API_KEY: str
    DEEPSEEK_API_KEY: str


# App initialization
app = FastAPI()

# Static files and templates
app.mount("/site", StaticFiles(directory="./frontend"), name="site")
app.mount("/static", StaticFiles(directory="./frontend/static"), name="static")
templates = Jinja2Templates(directory="./frontend")

# WebSocket manager
manager = WebSocketManager()

# Middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Constants
DOC_PATH = os.getenv("DOC_PATH", "./my-docs")

# Startup event


@app.get('/health')
async def root():
    return {
        "status": "ok",
        "version": "0.0.1",
    }

@app.on_event("startup")
def startup_event():
    os.makedirs("outputs", exist_ok=True)
    app.mount("/outputs", StaticFiles(directory="outputs"), name="outputs")
    os.makedirs(DOC_PATH, exist_ok=True)
    

# Routes

@app.post("/create-folder/{folder_name}")
async def create_folder(folder_name: str):
    """
    Create a new folder under the DOC_PATH directory.
    """
    folder_path = os.path.join(DOC_PATH, folder_name)
    
    # Check if folder already exists
    if os.path.exists(folder_path):
        raise HTTPException(status_code=400, detail=f"Folder '{folder_name}' already exists.")
    
    try:
        os.makedirs(folder_path, exist_ok=True)
        return {"status": "success", "message": f"Folder '{folder_name}' created successfully."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



@app.get("/list-folders")
async def list_folders():
    """
    List all folders under the DOC_PATH directory.
    """
    try:
        # Get all directories under DOC_PATH
        folders = [f for f in os.listdir(DOC_PATH) if os.path.isdir(os.path.join(DOC_PATH, f))]
        return {"status": "success", "folders": folders}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/")
async def read_root(request: Request):
    return templates.TemplateResponse("index.html", {"request": request, "report": None})


@app.get("/{folder_name}/files")
async def list_files(folder_name: str = None):
    """
    List files in a specific folder under the DOC_PATH directory.
    """
    folder_path = os.path.join(DOC_PATH, folder_name) if folder_name else DOC_PATH
    if not os.path.exists(folder_path):
        raise HTTPException(status_code=404, detail=f"Folder '{folder_name}' does not exist.")

    files = os.listdir(folder_path)
    return {"files": files}



@app.post("/api/multi_agents")
async def run_multi_agents():
    return await execute_multi_agents(manager)


@app.post("/upload/{folder_name}")
async def upload_file(folder_name: str, file: UploadFile = File(...)):
    """
    Upload a file to a specific folder under the DOC_PATH directory.
    """
    folder_path = os.path.join(DOC_PATH, folder_name)
    if not os.path.exists(folder_path):
        raise HTTPException(status_code=404, detail=f"Folder '{folder_name}' does not exist.")

    return await handle_file_upload(file, folder_name)

@app.delete("/folders/{folder_name}")
async def delete_folder(folder_name: str):
    """
    Delete an entire folder and all its contents.
    Args:
        folder_name (str): The name of the folder to delete.
    Returns:
        dict: A message indicating success or failure.
    """
    try:
        # Construct the full folder path
        folder_path = os.path.join(DOC_PATH, folder_name)
        
        # Check if the folder exists
        if not os.path.exists(folder_path):
            raise HTTPException(status_code=404, detail=f"Folder '{folder_name}' not found.")
        
        # Delete the folder and its contents
        shutil.rmtree(folder_path)
        return {"status": "success", "message": f"Folder '{folder_name}' deleted successfully."}
    
    except Exception as e:
        logger.error(f"Error deleting folder: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/files/{folder_name}/{filename}")
async def delete_file(folder_name: str, filename: str):
    """
    Delete a specific file in a folder.
    Args:
        folder_name (str): The name of the folder containing the file.
        filename (str): The name of the file to delete.
    Returns:
        dict: A message indicating success or failure.
    """
    try:
        # Construct the full file path
        file_path = os.path.join(DOC_PATH, folder_name, filename)
        
        # Check if the file exists
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail=f"File '{filename}' not found in folder '{folder_name}'.")
        
        # Delete the file
        os.remove(file_path)
        return {"status": "success", "message": f"File '{filename}' deleted successfully from folder '{folder_name}'."}
    
    except Exception as e:
        logger.error(f"Error deleting file: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# @app.delete("/files/{folder_name}/{filename}")
# async def delete_file(folder_name: str, filename: str):
#     """
#     Delete a specific file in a folder.
#     Args:
#         folder_name (str): The name of the folder containing the file.
#         filename (str): The name of the file to delete.
#     Returns:
#         dict: A message indicating success or failure.
#     """
#     try:
#         # Construct the full file path
#         file_path = os.path.join(DOC_PATH, folder_name, filename)
        
#         # Check if the file exists
#         if not os.path.exists(file_path):
#             raise HTTPException(status_code=404, detail=f"File '{filename}' not found in folder '{folder_name}'.")
        
#         # Delete the file
#         os.remove(file_path)
#         return {"status": "success", "message": f"File '{filename}' deleted successfully from folder '{folder_name}'."}
    
#     except Exception as e:
#         logger.error(f"Error deleting file: {e}")
#         raise HTTPException(status_code=500, detail=str(e))

@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await manager.connect(websocket)
    try:
        await handle_websocket_communication(websocket, manager)
    except WebSocketDisconnect:
        await manager.disconnect(websocket)


# Returns report in a docx format

S3_BUCKET_NAME = "your-s3-bucket-name"  # Replace with your bucket name
S3_REGION = "us-east-1"  # Replace with your bucket's region
s3_client = boto3.client("s3")

from fastapi import HTTPException, Response
import os
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)






# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# AWS S3 Configuration
S3_BUCKET_NAME = "gptresearcher"  # Replace with your bucket name
S3_REGION = "us-east-1"  # Replace with your bucket's region
s3_client = boto3.client("s3")  # Initialize S3 client

@app.get("/generate-report", response_class=JSONResponse)
async def generate_report(
    prompt: str,
    tone: Tone,
    report_source: ReportSource,
    report_type: ReportType,
    folder_name: str = None,
    tavily_api_key: str = None,
    open_ai_key: str = None
):
    """
    Generate a report, save it as a .docx file in AWS S3, and return a download link.
    """
    try:
        # Update environment variables if API keys are provided
        if tavily_api_key:
            os.environ["TAVILY_API_KEY"] = tavily_api_key
            logger.info("Tavily API key set")

        if open_ai_key:
            os.environ["OPENAI_API_KEY"] = open_ai_key
            logger.info("OpenAI API key set")

        # Determine the document path
        if folder_name:
            document_path = os.path.join(DOC_PATH, folder_name)
            if not os.path.exists(document_path):
                raise HTTPException(status_code=404, detail=f"Folder '{folder_name}' does not exist.")
        else:
            document_path = DOC_PATH

        # Get the list of documents in the folder
        document_urls = [os.path.join(document_path, f) for f in os.listdir(document_path) if os.path.isfile(os.path.join(document_path, f))]
        logger.info(f"Found {len(document_urls)} documents in {document_path}")

        # Run the research task
        logger.info(f"Running agent with prompt: {prompt}")
        report = await run_agent(
            task=prompt,
            report_type=report_type.value,
            report_source=report_source.value,
            tone=tone.value,
            websocket=None,
            headers=None,
            source_urls=[],
            document_urls=document_urls,
            query_domains=[],
            config_path="default"
        )

        # Check if report is valid
        if report is None:
            logger.error("run_agent returned None")
            raise HTTPException(status_code=500, detail="Failed to generate report: LLM response was empty")

        # Convert report to string if it’s not already
        if isinstance(report, bytes):
            report_content = report.decode('utf-8')
        elif isinstance(report, str):
            report_content = report
        else:
            logger.error(f"Unexpected report type: {type(report)}")
            raise ValueError(f"Report must be a string or bytes, got {type(report)}")

        # Create a .docx file
        doc = Document()
        doc.add_heading("Generated Report", level=1)
        doc.add_paragraph(report_content)

        # Save the document to a BytesIO stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_bytes = file_stream.getvalue()
        file_stream.close()

        # Generate a unique filename (e.g., using timestamp or UUID)
        from datetime import datetime
        filename = f"generated_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        # Upload the file to S3
        try:
            s3_client.put_object(
                Bucket=S3_BUCKET_NAME,
                Key=filename,
                Body=file_bytes,
                ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            logger.info(f"Uploaded {filename} to S3 bucket {S3_BUCKET_NAME}")
        except ClientError as e:
            logger.error(f"Failed to upload to S3: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to upload to S3: {str(e)}")

        # Generate a presigned URL for download (expires in 1 hour by default)
        try:
            download_url = s3_client.generate_presigned_url(
                "get_object",
                Params={"Bucket": S3_BUCKET_NAME, "Key": filename},
                ExpiresIn=3600  # URL expires in 1 hour (3600 seconds)
            )
            logger.info(f"Generated presigned URL: {download_url}")
        except ClientError as e:
            logger.error(f"Failed to generate presigned URL: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to generate download URL: {str(e)}")

        # Return the download link as JSON
        return JSONResponse(
            content={
                "status": "success",
                "message": "Report generated and uploaded to S3",
                "download_url": download_url
            }
        )

    except HTTPException as e:
        raise e
    except Exception as e:
        logger.exception(f"Error generating report: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating report: {str(e)}")


# @app.get("/generate-report", response_class=Response)
# async def generate_report(
#     prompt: str,
#     tone: Tone,
#     report_source: ReportSource,
#     report_type: ReportType,
#     folder_name: str = None,
#     tavily_api_key: str = None,
#     open_ai_key: str = None
# ):
#     """
#     Generate a report and return it as a byte array.
#     """
#     try:
#         # Update environment variables if API keys are provided
#         if tavily_api_key:
#             os.environ["TAVILY_API_KEY"] = tavily_api_key
#             logger.info("Tavily API key set")

#         if open_ai_key:
#             os.environ["OPENAI_API_KEY"] = open_ai_key
#             logger.info("OpenAI API key set")

#         # Determine the document path
#         if folder_name:
#             document_path = os.path.join(DOC_PATH, folder_name)
#             if not os.path.exists(document_path):
#                 raise HTTPException(status_code=404, detail=f"Folder '{folder_name}' does not exist.")
#         else:
#             document_path = DOC_PATH

#         # Get the list of documents in the folder
#         document_urls = [os.path.join(document_path, f) for f in os.listdir(document_path) if os.path.isfile(os.path.join(document_path, f))]
#         logger.info(f"Found {len(document_urls)} documents in {document_path}")

#         # Run the research task
#         logger.info(f"Running agent with prompt: {prompt}")
#         report = await run_agent(
#             task=prompt,
#             report_type=report_type.value,
#             report_source=report_source.value,
#             tone=tone.value,
#             websocket=None,
#             headers=None,
#             source_urls=[],
#             document_urls=document_urls,
#             query_domains=[],
#             config_path="default"
#         )

#         # Check if report is valid
#         if report is None:
#             logger.error("run_agent returned None")
#             raise HTTPException(status_code=500, detail="Failed to generate report: LLM response was empty")

#         # Convert report to bytes
#         if isinstance(report, str):
#             report_bytes = report.encode('utf-8')
#             logger.info("Report generated as string, converted to bytes")
#         elif isinstance(report, bytes):
#             report_bytes = report
#             logger.info("Report generated as bytes")
#         else:
#             logger.error(f"Unexpected report type: {type(report)}")
#             raise ValueError(f"Report must be a string or bytes, got {type(report)}")

#         # Return the report as a byte response
#         return Response(
#             content=report_bytes,
#             media_type="application/octet-stream",
#             headers={"Content-Disposition": "attachment; filename=report.txt"}
#         )

#     except HTTPException as e:
#         # Re-raise HTTP exceptions (e.g., 404 for folder not found)
#         raise e
#     except Exception as e:
#         # Log the error and return a detailed 500 response
#         logger.exception(f"Error generating report: {str(e)}")
#         raise HTTPException(status_code=500, detail=f"Error generating report: {str(e)}")


# @app.get("/generate-report")
# async def generate_report(
#     prompt: str,  # Query parameter for the task/prompt
#     tone: Tone,  # Query parameter for the tone
#     report_source: ReportSource,  # Query parameter for the report source
#     report_type: ReportType,  # Query parameter for the report type
#     folder_name: str = None,  # Optional query parameter for the folder name
#     tavily_api_key: str = None,  # Optional query parameter for the Tavily API key
#     open_ai_key: str = None  # Optional query parameter for the OpenAI API key
# ):
#     """
#     Generate a report and return it as a .docx file in bytes.
#     """
#     try:
#         # If the user provides a Tavily API key, update the environment variables
#         if tavily_api_key:
#             os.environ["TAVILY_API_KEY"] = tavily_api_key

#         if open_ai_key:
#             os.environ["OPENAI_API_KEY"] = open_ai_key

#         # Determine the document path
#         if folder_name:
#             document_path = os.path.join(DOC_PATH, folder_name)
#             if not os.path.exists(document_path):
#                 raise HTTPException(status_code=404, detail=f"Folder '{folder_name}' does not exist.")
#         else:
#             document_path = DOC_PATH

#         # Get the list of documents in the folder
#         document_urls = [os.path.join(document_path, f) for f in os.listdir(document_path) if os.path.isfile(os.path.join(document_path, f))]

#         # Run the research task synchronously
#         report = await run_agent(
#             task=prompt,
#             report_type=report_type.value,  # Pass the enum value
#             report_source=report_source.value,  # Pass the enum value
#             tone=tone.value,  # Pass the enum value
#             websocket=None,  # No WebSocket for this endpoint
#             headers=None,
#             source_urls=[],  # No specific URLs
#             document_urls=document_urls,  # Use documents from the specified folder
#             query_domains=[],  # Add query domains if needed
#             config_path="default"  # Add config path if needed
#         )

#         # Convert the report to a .docx file
#         doc = Document()
#         doc.add_heading("Generated Report", level=1)
#         doc.add_paragraph(report)  # Add the report content to the document

#         # Save the document to a BytesIO stream
#         file_stream = io.BytesIO()
#         doc.save(file_stream)
#         file_bytes = file_stream.getvalue()  # Get the bytes of the document

#         # Return the .docx file as bytes
#         return Response(
#             content=file_bytes,
#             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             headers={"Content-Disposition": f"attachment; filename=generated_report.docx"}
#         )

#     except Exception as e:
#         # Handle errors
#         raise HTTPException(status_code=500, detail=str(e))


# @app.get("/generate-report")
# async def generate_report(
#     prompt: str,  # Query parameter for the task/prompt
#     tone: Tone,  # Query parameter for the tone
#     report_source: ReportSource,  # Query parameter for the report source
#     report_type: ReportType,  # Query parameter for the report type
#     folder_name: str = None,  # Optional query parameter for the folder name
#     tavily_api_key: str = None,  # Optional query parameter for the Tavily API key
#     open_ai_key: str = None  # Optional query parameter for the OpenAI API key
# ):
#     """
#     Generate a report using query parameters.
#     """
#     try:
#         # If the user provides a Tavily API key, update the environment variables
#         if tavily_api_key:
#             os.environ["TAVILY_API_KEY"] = tavily_api_key

#         if open_ai_key:
#             os.environ["OPENAI_API_KEY"] = open_ai_key

#         # Determine the document path
#         if folder_name:
#             document_path = os.path.join(DOC_PATH, folder_name)
#             if not os.path.exists(document_path):
#                 raise HTTPException(status_code=404, detail=f"Folder '{folder_name}' does not exist.")
#         else:
#             document_path = DOC_PATH

#         # Get the list of documents in the folder
#         document_urls = [os.path.join(document_path, f) for f in os.listdir(document_path) if os.path.isfile(os.path.join(document_path, f))]

#         # Run the research task synchronously
#         report = await run_agent(
#             task=prompt,
#             report_type=report_type.value,  # Pass the enum value
#             report_source=report_source.value,  # Pass the enum value
#             tone=tone.value,  # Pass the enum value
#             websocket=None,  # No WebSocket for this endpoint
#             headers=None,
#             source_urls=[],  # No specific URLs
#             document_urls=document_urls,  # Use documents from the specified folder
#             query_domains=[],  # Add query domains if needed
#             config_path="default"  # Add config path if needed
#         )

#         # Return the generated report
#         return {"status": "success", "report": report}

#     except Exception as e:
#         # Handle errors
#         raise HTTPException(status_code=500, detail=str(e))

# {
#   "prompt": "What is the exams project?",
#   "tone": "Analytical (critical evaluation and detailed examination of data and theories)",
#   "report_source": "local",
#   "report_type": "detailed_report"
# }